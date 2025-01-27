import json
from docx import Document
from docx.shared import Pt


# Helper: Apply formatting to a paragraph or run
def apply_formatting(run, formatting):
    if "font_size" in formatting:
        run.font.size = Pt(formatting["font_size"])
    if "font_style" in formatting:
        if formatting["font_style"] == "bold":
            run.bold = True
        elif formatting["font_style"] == "italic":
            run.italic = True


# Helper: Replace placeholder, even across multiple runs
def replace_placeholder(doc, placeholder, content):
    for paragraph in doc.paragraphs:
        runs_text = "".join(run.text for run in paragraph.runs)
        if placeholder in runs_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(content)
            return True
    return False


# Process Personal Information Section
def process_personal_information(doc, personal_data):
    replace_placeholder(doc, "{{name}}", personal_data["name"].upper())
    replace_placeholder(doc, "{{desired_role}}", personal_data["desired_role"])
    contact_info = (
        f"Phone: {personal_data['contact_info']['phone']} | "
        f"Email: {personal_data['contact_info']['email']} | "
        f"LinkedIn: {personal_data['contact_info']['linkedin']} | "
        f"Location: {personal_data['contact_info']['location']}"
    )
    replace_placeholder(doc, "{{contact_info}}", contact_info)


# Process Summary Section
def process_summary(doc, summary_data):
    content = f"{summary_data['static_content']} {summary_data['dynamic_content']}"
    replace_placeholder(doc, "{{summary}}", content)


# Process Skills Section
def process_skills(doc, skills_data):
    skills_content = ""
    for group in skills_data["groups"]:
        skills_content += f"{group['group_name']}: " + ", ".join(group["skills_list"]) + "\n"
    replace_placeholder(doc, "{{skills}}", skills_content.strip())


# Process Experience Section
def process_experience(doc, experience_data):
    experience_content = ""
    for role in experience_data["roles"]:
        experience_content += f"{role['title']} - {role['company']} ({role['dates']})\n"
        experience_content += f"{role['location']}\n"
        for responsibility in role["responsibilities"]:
            experience_content += f"• {responsibility}\n"
        experience_content += "\n"
    replace_placeholder(doc, "{{experience_bullets}}", experience_content.strip())


# Process Education Section
def process_education(doc, education_data):
    education_content = ""
    for entry in education_data["entries"]:
        education_content += f"{entry['degree']}, {entry['focus']}\n"
        education_content += f"{entry['school']}\n\n"
    replace_placeholder(doc, "{{education}}", education_content.strip())


# Process Key Achievements Section
def process_key_achievements(doc, achievements_data):
    achievements_content = ""
    for entry in achievements_data["entries"]:
        achievements_content += f"{entry['title']}\n"
        achievements_content += f"• {entry['description']}\n\n"
    replace_placeholder(doc, "{{key_achievements}}", achievements_content.strip())


# Process Projects Section
def process_projects(doc, projects_data):
    projects_content = ""
    for project in projects_data["entries"]:
        projects_content += f"{project['title']} – {project['company']} ({project['dates']})\n"
        projects_content += f"Goal: {project['goal']}\n"
        projects_content += f"Responsibilities: {project['responsibilities']}\n"
        projects_content += f"Results: {project['results']}\n\n"
    replace_placeholder(doc, "{{projects}}", projects_content.strip())


# Main Rendering Function
def render_resume(json_path, template_path, output_path):
    with open(json_path, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)

    doc = Document(template_path)

    # Process all sections
    process_personal_information(doc, data["personal_information"])
    process_summary(doc, data["summary"])
    process_skills(doc, data["skills"])
    process_experience(doc, data["experience"])
    process_education(doc, data["education"])
    process_key_achievements(doc, data["key_achievements"])
    process_projects(doc, data["projects"])

    doc.save(output_path)
    print("Resume successfully generated!")


# Paths
json_path = "test_data/resume.json"  # Path to your JSON file
template_path = "test_templates/resume_template.docx"  # Path to your Word template
output_path = "output/final_resume.docx"  # Output file path

# Execute
render_resume(json_path, template_path, output_path)
