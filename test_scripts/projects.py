from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

# Helper: Apply spacing and formatting
def apply_paragraph_formatting(paragraph, font_size=11, bold=False, space_before=0, space_after=6):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(font_size)
    run.bold = bold

    # Adjust line spacing
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)

    # Ensure no extra line spacing
    paragraph_format.line_spacing = Pt(12)  # Set fixed line spacing

# Helper: Add a project block
def add_project_entry(doc, project):
    # Add the project title
    title_paragraph = doc.add_paragraph()
    title_text = f"{project['title']} – {project['company']} ({project['dates']})"
    title_run = title_paragraph.add_run(title_text)
    apply_paragraph_formatting(title_paragraph, font_size=12, bold=True, space_after=6)

    # Add the project details
    details = [
        f"Goal: {project['goal']}",
        f"Responsibilities: {project['responsibilities']}",
        f"Results: {project['results']}"
    ]
    for detail in details:
        detail_paragraph = doc.add_paragraph(detail)
        apply_paragraph_formatting(detail_paragraph, font_size=11, space_after=6)

# Main function to process projects
def process_projects(doc, projects_data):
    # Locate and replace the {{projects_entries}} placeholder
    for paragraph in doc.paragraphs:
        if "{{projects_entries}}" in paragraph.text:
            paragraph.text = ""  # Clear placeholder

            # Add each project
            for project in projects_data["entries"]:
                add_project_entry(doc, project)

            break  # Stop once placeholder is replaced

# Full script for rendering projects
def render_projects(json_path, template_path, output_path):
    import json
    with open(json_path, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)

    doc = Document(template_path)
    process_projects(doc, data["projects"])
    doc.save(output_path)

# Test file paths (update accordingly)
json_path = "test_data/projects.json"  # JSON with projects data
template_path = "test_templates/projects_template.docx"  # Word template
output_path = "output/test_projects.docx"  # Output file

# Run the script
render_projects(json_path, template_path, output_path)


