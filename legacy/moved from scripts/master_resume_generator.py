import logging
from docx import Document
from docx.shared import Pt

# ...existing code...

def generate_master_resume(data):
    """Generate the master resume document."""
    doc = Document()
    doc.add_heading(data['name'], level=1)
    doc.add_paragraph(data['summary'])

    for section in data['sections']:
        doc.add_heading(section['title'], level=2)
        for item in section['items']:
            p = doc.add_paragraph(item)
            run = p.add_run()
            run.font.size = Pt(12)

    doc.save('master_resume.docx')
    logging.info("Master resume generated successfully.")

# ...existing code...
