from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from formatting import apply_dynamic_formatting

def add_hyperlink(paragraph, url, text, formatting):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = paragraph.add_run(text)
    apply_dynamic_formatting(new_run, formatting)
    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

def add_section_header(document, section_name, formatting):
    """Add a section header to the document."""
    header = document.add_paragraph()
    header_run = header.add_run(section_name)
    apply_dynamic_formatting(header_run, formatting.get("section_header", {}))
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_section_content(document, section_data, formatting):
    """Add section content to the document."""
    if isinstance(section_data, list):
        for item in section_data:
            paragraph = document.add_paragraph()
            content_run = paragraph.add_run(str(item))
            apply_dynamic_formatting(content_run, formatting.get("content", {}))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif isinstance(section_data, dict):
        for key, value in section_data.items():
            paragraph = document.add_paragraph()
            sub_header_run = paragraph.add_run(f"{key}: ")
            apply_dynamic_formatting(sub_header_run, formatting.get("sub_header", {}))
            content_run = paragraph.add_run()
            apply_dynamic_formatting(content_run, formatting.get("content", {}))
            if isinstance(value, list):
                content_run.text = ", ".join(map(str, value))
            elif isinstance(value, dict):
                content_run.text = json.dumps(value, indent=4)
            else:
                content_run.text = str(value)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:
        paragraph = document.add_paragraph()
        content_run = paragraph.add_run(str(section_data))
        apply_dynamic_formatting(content_run, formatting.get("content", {}))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_personal_information(document, personal_info, formatting):
    """Add the personal information section to the document without a section header."""
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(personal_info["name"])
    apply_dynamic_formatting(name_run, formatting.get("name", {}))
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    role_paragraph = document.add_paragraph()
    role_run = role_paragraph.add_run(personal_info["desired_role"])
    apply_dynamic_formatting(role_run, formatting.get("desired_role", {}))
    role_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    contact_info = personal_info["contact_info"]
    contact_info_text = f"{contact_info['phone']} | {contact_info['email']} | {contact_info['location']} | "
    contact_paragraph = document.add_paragraph()
    contact_run = contact_paragraph.add_run(contact_info_text)
    apply_dynamic_formatting(contact_run, formatting.get("contact_info", {}))
    add_hyperlink(contact_paragraph, contact_info['linkedin'], contact_info['linkedin'], formatting.get("contact_info", {}))
    contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_skills_section(document, skills_data, formatting):
    """Add the skills section to the document."""
    add_section_header(document, "Skills", formatting)
    for group in skills_data["groups"]:
        group_paragraph = document.add_paragraph()
        group_run = group_paragraph.add_run(f"{group['group_name']}")
        apply_dynamic_formatting(group_run, formatting.get("sub_header", {}))
        group_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        skills_paragraph = document.add_paragraph()
        skills_run = skills_paragraph.add_run(", ".join(group["skills_list"]))
        apply_dynamic_formatting(skills_run, formatting.get("content", {}))
        skills_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_experience_section(document, experience_data, formatting):
    """Add the experience section to the document."""
    add_section_header(document, "Experience", formatting)
    for entry in experience_data:
        company_paragraph = document.add_paragraph()
        company_run = company_paragraph.add_run(f"{entry['company']} - {entry['title']} ({entry['dates']})")
        apply_dynamic_formatting(company_run, formatting.get("sub_header", {}))
        company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        location_paragraph = document.add_paragraph()
        location_run = location_paragraph.add_run(f"Location: {entry['location']}")
        apply_dynamic_formatting(location_run, formatting.get("content", {}))
        location_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for responsibility in entry["responsibilities"]:
            responsibility_paragraph = document.add_paragraph(style='List Bullet')
            responsibility_run = responsibility_paragraph.add_run(responsibility)
            apply_dynamic_formatting(responsibility_run, formatting.get("content", {}))
            responsibility_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_education_section(document, education_data, formatting):
    """Add the education section to the document."""
    add_section_header(document, "Education", formatting)
    for entry in education_data:
        education_paragraph = document.add_paragraph()
        education_run = education_paragraph.add_run(f"{entry['school']} - {entry['degree']}")
        apply_dynamic_formatting(education_run, formatting.get("sub_header", {}))
        education_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        location_focus_paragraph = document.add_paragraph()
        location_focus_run = location_focus_paragraph.add_run(f"Location: {entry['location']}, Focus: {entry['focus']}")
        apply_dynamic_formatting(location_focus_run, formatting.get("content", {}))
        location_focus_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_certifications_section(document, certifications_data, formatting):
    """Add the certifications section to the document."""
    add_section_header(document, "Certifications", formatting)
    for group in certifications_data:
        group_paragraph = document.add_paragraph()
        group_run = group_paragraph.add_run(f"{group['group_name']}")
        apply_dynamic_formatting(group_run, formatting.get("sub_header", {}))
        group_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for certification in group["certifications_list"]:
            certification_paragraph = document.add_paragraph(style='List Bullet')
            certification_run = certification_paragraph.add_run(certification)
            apply_dynamic_formatting(certification_run, formatting.get("content", {}))
            certification_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_key_achievements_section(document, key_achievements_data, formatting):
    """Add the key achievements section to the document."""
    add_section_header(document, "Key Achievements", formatting)
    for achievement in key_achievements_data:
        achievement_paragraph = document.add_paragraph()
        achievement_title_run = achievement_paragraph.add_run(achievement["title"])
        apply_dynamic_formatting(achievement_title_run, formatting.get("sub_header", {}))
        achievement_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        achievement_description_paragraph = document.add_paragraph()
        achievement_description_run = achievement_description_paragraph.add_run(achievement["description"])
        apply_dynamic_formatting(achievement_description_run, formatting.get("content", {}))
        achievement_description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_projects_section(document, projects_data, formatting):
    """Add the projects section to the document."""
    add_section_header(document, "Projects", formatting)
    for project in projects_data:
        project_paragraph = document.add_paragraph()
        project_title_run = project_paragraph.add_run(project["title"])
        apply_dynamic_formatting(project_title_run, formatting.get("sub_header", {}))
        project_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        project_details_paragraph = document.add_paragraph()
        project_details_run = project_details_paragraph.add_run(f"Company: {project['company']}, Dates: {project['dates']}")
        apply_dynamic_formatting(project_details_run, formatting.get("content", {}))
        project_details_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        project_goal_paragraph = document.add_paragraph()
        project_goal_run = project_goal_paragraph.add_run("Goal: ")
        apply_dynamic_formatting(project_goal_run, formatting.get("sub_header", {}))
        project_goal_run = project_goal_paragraph.add_run(project["goal"])
        apply_dynamic_formatting(project_goal_run, formatting.get("content", {}))
        project_goal_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        project_responsibilities_paragraph = document.add_paragraph()
        project_responsibilities_run = project_responsibilities_paragraph.add_run("Responsibilities: ")
        apply_dynamic_formatting(project_responsibilities_run, formatting.get("sub_header", {}))
        project_responsibilities_run = project_responsibilities_paragraph.add_run(project["responsibilities"])
        apply_dynamic_formatting(project_responsibilities_run, formatting.get("content", {}))
        project_responsibilities_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        project_results_paragraph = document.add_paragraph()
        project_results_run = project_results_paragraph.add_run("Results: ")
        apply_dynamic_formatting(project_results_run, formatting.get("sub_header", {}))
        project_results_run = project_results_paragraph.add_run(project["results"])
        apply_dynamic_formatting(project_results_run, formatting.get("content", {}))
        project_results_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Debug statement to confirm the file is being executed
print("sections.py loaded successfully")
