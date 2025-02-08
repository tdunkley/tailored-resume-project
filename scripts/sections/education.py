from docx import Document
from docx.shared import Pt
import json

def generate_education_section(output_path):
    # Load data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)

    # Extract education data from resume.json
    education_info = resume.get("education", {}).get("entries", [])
    formatting = resume.get("education", {}).get("formatting", {})

    # Check if education_info is loaded correctly
    if not education_info:
        print("No education data available.")
        return

    # Create a new Word Document
    doc = Document()

    # Add Education Section Header (all caps, bold, size 12)
    education_header = doc.add_paragraph()
    header_run = education_header.add_run("EDUCATION")
    header_run.bold = True
    header_run.font.size = Pt(12)
    education_header.paragraph_format.space_before = Pt(6)
    education_header.paragraph_format.space_after = Pt(6)

    # Loop through each education entry
    for entry in education_info:
        # Add Degree & Major (sub-header) on the same line
        education_paragraph = doc.add_paragraph()
        school_degree_run = education_paragraph.add_run(f"{entry.get('degree', 'N/A')} in {entry.get('focus', 'N/A')} — {entry.get('school', 'N/A')}")
        school_degree_run.font.size = Pt(formatting.get("font_size", 11))
        school_degree_run.bold = True
        education_paragraph.paragraph_format.space_before = Pt(6)
        education_paragraph.paragraph_format.space_after = Pt(6)

        # Add Location (regular, size 11)
        location_paragraph = doc.add_paragraph()
        location_run = location_paragraph.add_run(f"Location: {entry.get('location', 'N/A')}")
        location_run.font.size = Pt(formatting.get("font_size", 11))
        location_paragraph.paragraph_format.space_before = Pt(6)
        location_paragraph.paragraph_format.space_after = Pt(6)

    # Save the document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")

# Example usage
output_path = "output/education_section.docx"
generate_education_section(output_path)
