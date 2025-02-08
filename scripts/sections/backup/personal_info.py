import os
import json
from docx import Document
from docx.shared import Pt, RGBColor

def load_resume_data():
    """
    Dynamically loads the resume.json file from the resume_project folder.
    
    Returns:
        dict: Loaded resume data.
    """
    # Get the absolute path to the 'resume_project' folder, no matter where the script is run from
    script_dir = os.path.dirname(os.path.realpath(__file__))  # Directory where the script is
    project_dir = os.path.abspath(os.path.join(script_dir, '..', '..', '..', 'resume_project'))  # Navigate to 'resume_project'

    # Now, calculate the correct path to resume.json
    json_file_path = os.path.join(project_dir, 'resume.json')

    print(f"Looking for resume.json at: {json_file_path}")  # Debugging output

    # Check if resume.json exists
    if not os.path.exists(json_file_path):
        raise FileNotFoundError(f"Could not find resume.json at {json_file_path}")
    
    # Load and return the resume data
    with open(json_file_path, 'r') as f:
        return json.load(f)


def generate_personal_info_section(output_path):
    """
    Generate the personal info section of the resume.

    Args:
        output_path (str): The path where the personal info section will be saved.
    """
    # Load resume data dynamically
    resume_data = load_resume_data()

    # Extract personal information from resume data
    personal_info = resume_data["personal_information"]

    # Create a new Word Document
    doc = Document()

    # Add Name
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(personal_info["name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(14)

    # Add Desired Role
    role_paragraph = doc.add_paragraph()
    role_run = role_paragraph.add_run(personal_info["desired_role"].title())
    role_run.italic = True
    role_run.font.size = Pt(12)

    # Add Contact Information
    contact_paragraph = doc.add_paragraph()
    contact_info = personal_info["contact_info"]
    contact_order = ["phone", "email", "linkedin", "location"]

    for index, key in enumerate(contact_order):
        value = contact_info.get(key, "")
        if key == "linkedin" and value:
            # Add clickable LinkedIn URL
            hyperlink = contact_paragraph.add_run(value)
            hyperlink.font.size = Pt(11)
            hyperlink.font.color.rgb = RGBColor(0, 0, 255)
            hyperlink.underline = True
            if index != len(contact_order) - 1:
                contact_paragraph.add_run(" | ").font.size = Pt(11)
        elif value:
            contact_paragraph.add_run(f"{value}")
            if index != len(contact_order) - 1:
                contact_paragraph.add_run(" | ").font.size = Pt(11)

    # Ensure the output directory exists
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the document
    doc.save(output_path)
    print(f"Personal info section generated: {output_path}")

# Example usage
if __name__ == "__main__":
    # Path should be relative to the root, not within 'scripts'
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'output', 'personal_info_section.docx')
    generate_personal_info_section(output_path)




