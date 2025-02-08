import os
from docx import Document
from docx.shared import Pt
import json

def load_resume_data():
    """
    Dynamically loads the resume.json file from the resume_project folder.
    
    Returns:
        dict: Loaded resume data.
    """
    # Get the absolute path to the 'resume_project' folder, no matter where the script is run from
    script_dir = os.path.dirname(os.path.realpath(__file__))  # Directory where the script is
    project_dir = os.path.abspath(os.path.join(script_dir, '..', '..', '..', 'resume_project'))  # Navigate to 'resume_project'

    # Now, calculate the correct path to resume.json
    json_file_path = os.path.join(project_dir, 'resume.json')

    print(f"Looking for resume.json at: {json_file_path}")  # Debugging output

    # Check if resume.json exists
    if not os.path.exists(json_file_path):
        raise FileNotFoundError(f"Could not find resume.json at {json_file_path}")
    
    # Load and return the resume data
    with open(json_file_path, 'r') as f:
        return json.load(f)


def generate_summary_section(output_path):
    """
    Generate the summary section of the resume.

    Args:
        output_path (str): The path where the summary section will be saved.
    """
    # Load resume data dynamically
    resume_data = load_resume_data()

    # Create a new Word Document for the Summary Section
    doc = Document()

    # Add Summary Section Header
    summary_header = doc.add_paragraph()
    header_run = summary_header.add_run("SUMMARY")
    header_run.bold = True
    header_run.font.size = Pt(12)
    summary_header.paragraph_format.space_before = Pt(6)
    summary_header.paragraph_format.space_after = Pt(6)

    # Add Summary Content (Static + Dynamic Content)
    summary_content = doc.add_paragraph()
    content_run = summary_content.add_run(
        "Visionary Data and Analytics Executive with over 15 years of experience. "
        "Specializes in leveraging data to drive decision-making, enhance operations, and achieve measurable business outcomes."
    )
    content_run.italic = True
    content_run.font.size = Pt(12)
    summary_content.paragraph_format.space_before = Pt(6)
    summary_content.paragraph_format.space_after = Pt(6)

    # Ensure the output directory exists
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the Word document
    doc.save(output_path)
    print(f"Summary section generated: {output_path}")

# Example usage
if __name__ == "__main__":
    # Dynamically calculate the output path relative to the root directory
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'output', 'summary_section.docx')
    generate_summary_section(output_path)

