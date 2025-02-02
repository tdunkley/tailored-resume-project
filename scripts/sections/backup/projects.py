from docx import Document
from docx.shared import Pt
import json

def generate_projects_section(output_path):
    # Load data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)

    # Extract projects data from resume.json
    projects_info = resume.get("projects", {}).get("entries", [])
    formatting = resume.get("projects", {}).get("formatting", {})

    # Check if projects_info is loaded correctly
    if not projects_info:
        print("No projects data available.")
        return

    # Create a new Word Document
    doc = Document()

    # Add Projects Section Header (all caps, bold, size 12)
    projects_header = doc.add_paragraph()
    header_run = projects_header.add_run("PROJECTS")
    header_run.bold = True
    header_run.font.size = Pt(12)
    projects_header.paragraph_format.space_before = Pt(6)
    projects_header.paragraph_format.space_after = Pt(6)

    # Loop through each project entry
    for entry in projects_info:
        # Add Title and Company Name (bold, size 12)
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(f"{entry.get('title', 'N/A')} — {entry.get('company', 'N/A')}")
        title_run.font.size = Pt(formatting.get("title", {}).get("font_size", 12))
        title_run.bold = True
        title_paragraph.paragraph_format.space_before = Pt(6)
        title_paragraph.paragraph_format.space_after = Pt(6)

        # Add Goal (regular, size 11, bold for "Goal" only)
        goal_paragraph = doc.add_paragraph()
        goal_run = goal_paragraph.add_run("Goal: ")
        goal_run.bold = True  # Bold the "Goal"
        goal_run.font.size = Pt(formatting.get("details", {}).get("font_size", 11))
        goal_paragraph.add_run(f" {entry.get('goal', 'N/A')}")
        goal_paragraph.paragraph_format.space_before = Pt(6)
        goal_paragraph.paragraph_format.space_after = Pt(6)

        # Add Responsibilities (regular, size 11, bold for "Responsibilities" only)
        responsibilities_paragraph = doc.add_paragraph()
        responsibilities_run = responsibilities_paragraph.add_run("Responsibilities: ")
        responsibilities_run.bold = True  # Bold the "Responsibilities"
        responsibilities_run.font.size = Pt(formatting.get("details", {}).get("font_size", 11))
        responsibilities_paragraph.add_run(f" {entry.get('responsibilities', 'N/A')}")
        responsibilities_paragraph.paragraph_format.space_before = Pt(6)
        responsibilities_paragraph.paragraph_format.space_after = Pt(6)

        # Add Results (regular, size 11, bold for "Results" only)
        results_paragraph = doc.add_paragraph()
        results_run = results_paragraph.add_run("Results: ")
        results_run.bold = True  # Bold the "Results"
        results_run.font.size = Pt(formatting.get("details", {}).get("font_size", 11))
        results_paragraph.add_run(f" {entry.get('results', 'N/A')}")
        results_paragraph.paragraph_format.space_before = Pt(6)
        results_paragraph.paragraph_format.space_after = Pt(6)

    # Save the document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")

# Example usage
output_path = "output/projects_section.docx"
generate_projects_section(output_path)


