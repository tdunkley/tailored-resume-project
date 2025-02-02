from docx import Document
from docx.shared import Pt
import json

def generate_experience_section(output_path):
    # Load data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)

    # Extract experience data from resume.json
    experience_info = resume.get("experience", {}).get("roles", [])
    formatting = resume.get("experience", {}).get("formatting", {})

    # Check if experience_info is loaded correctly
    print("Experience Section Data:")
    print(experience_info)

    if not experience_info:
        print("No experience data available.")
        return

    # Create a new Word Document
    doc = Document()

    # Add Experience Section Header (all caps, bold, size 12)
    experience_header = doc.add_paragraph()
    header_run = experience_header.add_run("EXPERIENCE")
    header_run.bold = True
    header_run.font.size = Pt(12)
    experience_header.paragraph_format.space_before = Pt(6)
    experience_header.paragraph_format.space_after = Pt(6)

    # Loop through each job in experience_info
    for job in experience_info:
        # Add Company, Title, and Dates on the same line
        job_paragraph = doc.add_paragraph()
        job_run = job_paragraph.add_run(f"{job.get('company', 'N/A')} — {job.get('title', 'N/A')} ({job.get('dates', 'N/A')})")
        job_run.bold = True
        job_run.font.size = Pt(formatting.get("title", {}).get("font_size", 12))
        job_paragraph.paragraph_format.space_before = Pt(6)
        job_paragraph.paragraph_format.space_after = Pt(6)

        # Add Location (regular, size 11)
        location_paragraph = doc.add_paragraph()
        location_run = location_paragraph.add_run(f"Location: {job.get('location', 'N/A')}")
        location_run.font.size = Pt(formatting.get("location", {}).get("font_size", 11))
        location_paragraph.paragraph_format.space_before = Pt(6)
        location_paragraph.paragraph_format.space_after = Pt(6)

        # Add Responsibilities (Bulleted list with correct font size and bullet style)
        responsibilities = job.get("responsibilities", [])
        for responsibility in responsibilities:
            bullet_paragraph = doc.add_paragraph(responsibility, style="List Bullet")
            bullet_paragraph.runs[0].font.size = Pt(formatting.get("responsibilities", {}).get("font_size", 11))
            bullet_paragraph.paragraph_format.space_before = Pt(formatting.get("responsibilities", {}).get("line_spacing", {}).get("before", 6))
            bullet_paragraph.paragraph_format.space_after = Pt(formatting.get("responsibilities", {}).get("line_spacing", {}).get("after", 6))

    # Save the document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")

# Example usage
output_path = "output/experience_section.docx"
generate_experience_section(output_path)
