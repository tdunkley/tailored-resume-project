from docx import Document
from docx.shared import Pt
import json

def generate_key_achievements_section(output_path):
    # Load data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)

    # Extract key achievements data from resume.json
    key_achievements_info = resume.get("key_achievements", {}).get("entries", [])
    formatting = resume.get("key_achievements", {}).get("formatting", {})

    # Check if key_achievements_info is loaded correctly
    if not key_achievements_info:
        print("No key achievements data available.")
        return

    # Create a new Word Document
    doc = Document()

    # Add Key Achievements Section Header (all caps, bold, size 12)
    key_achievements_header = doc.add_paragraph()
    header_run = key_achievements_header.add_run("KEY ACHIEVEMENTS")
    header_run.bold = True
    header_run.font.size = Pt(12)
    key_achievements_header.paragraph_format.space_before = Pt(6)
    key_achievements_header.paragraph_format.space_after = Pt(6)

    # Loop through each key achievement entry
    for entry in key_achievements_info:
        # Add Title (bold, size 12)
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(entry.get('title', 'N/A'))
        title_run.font.size = Pt(formatting.get("title", {}).get("font_size", 12))
        title_run.bold = True
        title_paragraph.paragraph_format.space_before = Pt(6)
        title_paragraph.paragraph_format.space_after = Pt(6)

        # Add Description (regular, size 11)
        description_paragraph = doc.add_paragraph()
        description_run = description_paragraph.add_run(entry.get('description', 'N/A'))
        description_run.font.size = Pt(formatting.get("description", {}).get("font_size", 11))
        description_paragraph.paragraph_format.space_before = Pt(6)
        description_paragraph.paragraph_format.space_after = Pt(6)

    # Save the document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")

# Example usage
output_path = "output/key_achievements_section.docx"
generate_key_achievements_section(output_path)
