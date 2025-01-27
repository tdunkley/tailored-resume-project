import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Helper: Apply formatting to a paragraph
def apply_formatting(paragraph, formatting):
    if formatting.get("font_size"):
        for run in paragraph.runs:
            run.font.size = Pt(formatting["font_size"])
    if formatting.get("font_style"):
        for run in paragraph.runs:
            if formatting["font_style"] == "bold":
                run.bold = True
            elif formatting["font_style"] == "italic":
                run.italic = True
    if formatting.get("alignment"):
        paragraph.alignment = formatting["alignment"]

# Helper: Add a bullet point with consistent formatting
def add_bullet_paragraph(document, text):
    try:
        bullet_paragraph = document.add_paragraph(style="List Bullet")
    except KeyError:
        bullet_paragraph = document.add_paragraph()  # Fallback to default style
    run = bullet_paragraph.add_run(text)
    run.font.size = Pt(11)  # Consistent bullet size
    bullet_paragraph.paragraph_format.space_after = Pt(6)  # Adjust spacing
    return bullet_paragraph

# Process experience section with consistent formatting
def process_experience(doc, experience_data):
    # Replace the experience header
    for paragraph in doc.paragraphs:
        if "{{experience_header}}" in paragraph.text:
            paragraph.text = "Professional Experience"
            apply_formatting(paragraph, {"font_size": 14, "font_style": "bold", "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT})
            break

    # Locate the placeholder and insert roles dynamically
    for paragraph in doc.paragraphs:
        if "{{experience_bullets}}" in paragraph.text:
            paragraph.text = ""  # Clear placeholder

            for role in experience_data["roles"]:
                # Add role title, company, and dates
                role_paragraph = doc.add_paragraph()
                role_text = f"{role['title']} - {role['company']} ({role['dates']})"
                role_run = role_paragraph.add_run(role_text)
                role_run.bold = True
                role_run.font.size = Pt(12)
                role_paragraph.paragraph_format.space_after = Pt(4)  # Adjust spacing

                # Add location
                location_paragraph = doc.add_paragraph()
                location_paragraph.add_run(role["location"]).font.size = Pt(11)
                location_paragraph.paragraph_format.space_after = Pt(2)  # Adjust spacing

                # Add responsibilities as bullets
                for responsibility in role["responsibilities"]:
                    add_bullet_paragraph(doc, responsibility)
            break  # End once the placeholder is replaced

# Render resume for testing
def render_resume_experience_test(json_path, template_path, output_path):
    with open(json_path, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)

    doc = Document(template_path)
    process_experience(doc, data["experience"])
    doc.save(output_path)

# Paths for testing
json_path = "test_data/experience.json"  # Ensure this file contains valid experience data
template_path = "test_templates/experience_template.docx"  # Ensure it contains {{experience_header}} and {{experience_bullets}}
output_path = "output/test_experience.docx"

# Run the script
render_resume_experience_test(json_path, template_path, output_path)













