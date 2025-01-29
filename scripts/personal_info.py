from docx import Document
from docx.shared import Pt, RGBColor
import json

def generate_personal_info_section(output_path):
    # Load the resume data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)
    
    # Create a new Word Document
    doc = Document()
    
    # Extract personal information
    personal_info = resume["personal_information"]

    # Add Name
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(personal_info["name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(14)

    # Add Desired Role
    role_paragraph = doc.add_paragraph()
    role_run = role_paragraph.add_run(personal_info["desired_role"].title())
    role_run.italic = True
    role_run.font.size = Pt(12)

    # Add Contact Information
    contact_paragraph = doc.add_paragraph()
    contact_info = personal_info["contact_info"]
    contact_order = ["phone", "email", "linkedin", "location"]

    for index, key in enumerate(contact_order):
        value = contact_info.get(key, "")
        if key == "linkedin" and value:
            # Add clickable LinkedIn URL
            hyperlink = contact_paragraph.add_run(value)
            hyperlink.font.size = Pt(11)
            hyperlink.font.color.rgb = RGBColor(0, 0, 255)
            hyperlink.underline = True
            if index != len(contact_order) - 1:
                contact_paragraph.add_run(" | ").font.size = Pt(11)
        elif value:
            contact_paragraph.add_run(f"{value}")
            if index != len(contact_order) - 1:
                contact_paragraph.add_run(" | ").font.size = Pt(11)

    # Save the document
    doc.save(output_path)
    print(f"Personal info section generated: {output_path}")

# Example usage
if __name__ == "__main__":
    output_path = "output/personal_info_section.docx"
    generate_personal_info_section(output_path)
