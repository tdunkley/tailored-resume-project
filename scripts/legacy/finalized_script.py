from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def render_resume(json_data, output_path):
    # Create a new Word document
    doc = Document()

    # Helper function for adding headings
    def add_heading(text, level=1):
        heading = doc.add_heading(level=level)
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Helper function for adding paragraphs
    def add_paragraph(text, bold=False, italic=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_after = Pt(8)

    # Add Personal Information
    personal_info = json_data.get("personal_information", {})
    contact_info = personal_info.get("contact_info", {})
    add_heading(personal_info.get("name", "NAME").upper())
    contact_details = (
        f"Phone: {contact_info.get('phone', '')} | "
        f"Email: {contact_info.get('email', '')} | "
        f"Location: {contact_info.get('location', '')} | "
        f"LinkedIn: {contact_info.get('linkedin', '')}"
    )
    add_paragraph(contact_details)

    # Add Summary
    add_heading("Summary")
    summary = json_data.get("summary", {}).get("content", "")
    add_paragraph(summary)

    # Add Skills
    add_heading("Skills")
    skills = json_data.get("skills", {})
    static_skills = ", ".join(skills.get("static", []))
    dynamic_skills = ", ".join(skills.get("dynamic", []))
    add_paragraph(f"Technical Skills: {static_skills}")
    add_paragraph(f"Dynamic Skills: {dynamic_skills}")

    # Add Experience
    add_heading("Experience")
    experiences = json_data.get("experience", {}).get("roles", [])
    for role in experiences:
        add_paragraph(f"{role.get('company')} | {role.get('title')} | {role.get('location')} | {role.get('dates')}", bold=True)
        for responsibility in role.get("responsibilities", []):
            add_paragraph(f"• {responsibility}")

    # Add Education
    add_heading("Education")
    education = json_data.get("education", {}).get("entries", [])
    for entry in education:
        add_paragraph(f"{entry.get('school')} | {entry.get('degree')} | {entry.get('location')}")

    # Add Key Achievements
    add_heading("Key Achievements")
    achievements = json_data.get("key_achievements", {}).get("entries", [])
    for achievement in achievements:
        add_paragraph(f"{achievement.get('title')}", bold=True)
        add_paragraph(achievement.get("description", ""))

    # Add Projects
    add_heading("Projects")
    projects = json_data.get("projects", {}).get("entries", [])
    for project in projects:
        add_paragraph(f"{project.get('title')}", bold=True)
        add_paragraph(f"Goal: {project.get('goal')}")
        add_paragraph(f"Responsibilities: {project.get('responsibilities')}")
        add_paragraph(f"Results: {project.get('results')}")

    # Save the document
    doc.save(output_path)
    print(f"Resume saved to {output_path}")

# Example usage
if __name__ == "__main__":
    # JSON data (replace with the actual JSON content)
    import json

    # Use the finalized JSON schema provided earlier
    json_data = {
        "personal_information": {
            "name": "Troy D. Dunkley",
            "contact_info": {
                "phone": "+1-770-401-6527",
                "email": "tdunkley@gmail.com",
                "location": "Atlanta, GA",
                "linkedin": "https://www.linkedin.com/in/troy-d-dunkley"
            }
        },
        # Add other sections as per the JSON schema...
    }

    output_file = "tailored_resume.docx"
    render_resume(json_data, output_file)
