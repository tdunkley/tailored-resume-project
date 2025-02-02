
from docx import Document
from docx.shared import Pt

def process_skills(doc, skills_data):
    # Add Section Header ("SKILLS")
    header_paragraph = doc.add_paragraph()
    header_run = header_paragraph.add_run("SKILLS")
    header_run.bold = True
    header_run.font.size = Pt(12)

    # Loop through the skill groups to apply consistent formatting with proper spacing
    for group in skills_data["groups"]:
        # Add Group Name as Subheader (bold, font size 12)
        group_name_paragraph = doc.add_paragraph()
        group_name_run = group_name_paragraph.add_run(group["group_name"])
        group_name_run.bold = True
        group_name_run.font.size = Pt(12)

        # Apply proper spacing after the group name
        group_name_paragraph.paragraph_format.space_after = Pt(6)

        # Add Skills List below the group name (font size 11)
        skills_paragraph = doc.add_paragraph()
        skills_run = skills_paragraph.add_run(", ".join(group["skills_list"]))
        skills_run.font.size = Pt(11)

        # Apply spacing after the skills list to ensure separation from next group
        skills_paragraph.paragraph_format.space_after = Pt(6)

# Example Usage:
if __name__ == "__main__":
    # Create a new document
    doc = Document()

    # Define the skills data (to be replaced by your actual data)
    skills_data = {
        "groups": [
            {"group_name": "Technical Skills", "skills_list": ["Python", "SQL", "Power BI", "AWS"]},
            {"group_name": "Soft Skills", "skills_list": ["Data Governance", "Predictive Analytics"]}
        ]
    }

    # Process the Skills section
    process_skills(doc, skills_data)

    # Save the document
    doc.save("formatted_skills_section.docx")
