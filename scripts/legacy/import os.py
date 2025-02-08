from docx import Document

def generate_personal_info_section(output_path):
    doc = Document()

    # Add Name
    doc.add_paragraph("Troy D. Dunkley", style='Heading 1')

    # Add Desired Role
    doc.add_paragraph("Data Analytics Executive", style='Heading 2')

    # Add Contact Information
    doc.add_paragraph("Phone: +1-770-401-6527")
    doc.add_paragraph("Email: tdunkley@gmail.com")
    doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/troy-d-dunkley")
    doc.add_paragraph("Location: Atlanta, GA")

    # Save document with error handling
    try:
        doc.save(output_path)
        print(f"Document saved successfully at {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

# Example usage
output_path = "C:/Users/troy_/Downloads/Tailored Resume Process/resume_project/output/full_resume.docx"
generate_personal_info_section(output_path)
