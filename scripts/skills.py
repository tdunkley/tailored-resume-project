from docx import Document
from docx.shared import Pt
import json

def generate_skills_section(output_path):
    # Load data from resume.json
    with open('resume.json', 'r') as f:
        resume = json.load(f)

    # Extract skills section info from resume.json
    skills_info = resume["skills"]

    # Create a new Word Document
    doc = Document()

    # Add Skills Section Header (all caps, bold, size 12)
    skills_header = doc.add_paragraph()
    header_run = skills_header.add_run("SKILLS")
    header_run.bold = True
    header_run.font.size = Pt(12)
    skills_header.paragraph_format.space_before = Pt(6)
    skills_header.paragraph_format.space_after = Pt(6)

    # Add Skill Groups and Skills List from resume.json
    for group in skills_info["groups"]:
        # Add Group Name (Title Case, Bold)
        group_paragraph = doc.add_paragraph()
        group_run = group_paragraph.add_run(group["group_name"].title())
        group_run.bold = True
        group_run.font.size = Pt(12)
        group_paragraph.paragraph_format.space_before = Pt(6)
        group_paragraph.paragraph_format.space_after = Pt(6)
        
        # Add Skills List (Sentence Case, Size 11)
        skills_paragraph = doc.add_paragraph()
        skills_paragraph.add_run(", ".join(group["skills_list"])).font.size = Pt(11)

    # Save the document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")

# Example usage
output_path = "output/skills_section.docx"
generate_skills_section(output_path)
