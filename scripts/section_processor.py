<<<<<<< HEAD
from docx import Document
import os

def process_sections(section_name, resume_data, output_dir):
    section_data = resume_data.get(section_name)
    if not section_data:
        print(f"Error: Section '{section_name}' not found.")
        return False

    print(f"Inline Preview for Section '{section_name}':")
    for key, value in section_data.items():
        print(f"{key}: {value}")
    print("-" * 50)

    # Create a new document from scratch
    document = Document()
    document.add_heading(section_name.upper(), level=1)
    for key, value in section_data.items():
        document.add_heading(key, level=2)
        document.add_paragraph(str(value))

    output_file = os.path.join(output_dir, f"{section_name}.docx")
    document.save(output_file)
    print(f"Document saved to {output_file}")
    return True
=======
import os
import sys
import logging
from core import get_paths
from tracker import CentralizedTracker
from validation_engine import validate_sectional_rules

# Dynamically add the scripts directory to the Python path
current_dir = os.path.dirname(os.path.abspath(__file__))
scripts_dir = os.path.abspath(os.path.join(current_dir, ".."))
if scripts_dir not in sys.path:
    sys.path.append(scripts_dir)

def process_sections (section_name, resume_data, tracker, paths):
    """
    Processes a single section from the resume JSON file.
    
    Parameters:
    - section_name (str): Name of the section to process.
    - resume_data (dict): Parsed resume JSON data.
    - tracker (CentralizedTracker): Tracker instance for logging and monitoring.
    - paths (dict): Dictionary containing key paths (e.g., output directory).
    """
    try:
        tracker.track_section_start(section_name)
        logger = logging.getLogger(f"SectionProcessor-{section_name}")
        logger.info(f"Processing section: {section_name}...")

        # Retrieve section data
        section_data = resume_data.get("sections", {}).get(section_name, None)
        if not section_data:
            raise ValueError(f"Section '{section_name}' not found in the resume data.")

        # Validate the section
        logger.info(f"Validating section: {section_name}...")
        validation_errors = validate_sectional_rules(section_name, section_data)
        if validation_errors:
            tracker.track_error(section_name, f"Validation errors: {validation_errors}")
            raise ValueError(f"Validation failed for section '{section_name}': {validation_errors}")

        # Format the section output (example formatting logic)
        logger.info(f"Formatting section: {section_name}...")
        formatted_output = f"{section_name.upper()}\n\n"
        for key, value in section_data.items():
            formatted_output += f"{key.title()}: {value}\n"

        # Write the formatted section to the output directory
        output_dir = paths.get("output_dir")
        os.makedirs(output_dir, exist_ok=True)
        output_file = os.path.join(output_dir, f"{section_name}.txt")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(formatted_output)

        # Log and track successful processing
        logger.info(f"Section '{section_name}' written to {output_file} successfully.")
        tracker.track_section_success(section_name)

    except Exception as e:
        logger.error(f"An error occurred while processing section '{section_name}': {e}", exc_info=True)
        tracker.track_error(section_name, str(e))
>>>>>>> 011eb20403b89d07c8ae11e5b4bc3094a262060f
