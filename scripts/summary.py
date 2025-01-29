from docx import Document
from docx.shared import Pt

def generate_summary_section(output_path):
    # Create a new Word Document for the Summary Section
    doc = Document()

    # Add Summary Section Header
    summary_header = doc.add_paragraph()
    header_run = summary_header.add_run("SUMMARY")
    header_run.bold = True
    header_run.font.size = Pt(12)
    summary_header.paragraph_format.space_before = Pt(6)
    summary_header.paragraph_format.space_after = Pt(6)

    # Add Summary Content (Static + Dynamic Content)
    summary_content = doc.add_paragraph()
    content_run = summary_content.add_run(
        "Visionary Data and Analytics Executive with over 15 years of experience. "
        "Specializes in leveraging data to drive decision-making, enhance operations, and achieve measurable business outcomes."
    )
    content_run.italic = True
    content_run.font.size = Pt(12)
    summary_content.paragraph_format.space_before = Pt(6)
    summary_content.paragraph_format.space_after = Pt(6)

    # Save the Word document
    doc.save(output_path)

# Example usage
if __name__ == "__main__":
    generate_summary_section("output/summary_section.docx")
