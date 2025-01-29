from docx import Document
from docx.shared import Pt, RGBColor

# Define sample data including LinkedIn as required
contact_info = {
    "phone": "+1-770-401-6527",
    "email": "tdunkley@gmail.com",
    "linkedin": "https://www.linkedin.com/in/troy-d-dunkley",
    "location": "Atlanta, GA"
}

# Create a new Document
doc = Document()

# Add Name
name_paragraph = doc.add_paragraph()
name_run = name_paragraph.add_run("TROY D. DUNKLEY")
name_run.bold = True
name_run.font.size = Pt(14)

# Add Desired Role
role_paragraph = doc.add_paragraph()
role_run = role_paragraph.add_run("Senior Data Architect")
role_run.italic = True
role_run.font.size = Pt(12)

# Add Contact Info with proper formatting and clickable LinkedIn URL
contact_paragraph = doc.add_paragraph()
contact_order = ["phone", "email", "linkedin", "location"]

# Add contact information with proper pipe separators
for index, key in enumerate(contact_order):
    value = contact_info.get(key, "")
    if key == "linkedin" and value:
        # Make the LinkedIn URL clickable
        contact_paragraph.add_run("LinkedIn: ").font.size = Pt(11)
        hyperlink = contact_paragraph.add_run(value)
        hyperlink.font.size = Pt(11)
        hyperlink.font.color.rgb = RGBColor(0, 0, 255)
        hyperlink.underline = True
        hyperlink.hyperlink = value
        contact_paragraph.add_run(" | ").font.size = Pt(11)  # Add pipe separator after LinkedIn
    elif value:
        # Add other contact fields with separators
        if index != len(contact_order) - 1:  # Avoid trailing pipe
            contact_paragraph.add_run(f"{key.title()}: {value} | ").font.size = Pt(11)
        else:
            contact_paragraph.add_run(f"{key.title()}: {value}").font.size = Pt(11)

# Save the document with the clickable LinkedIn URL and proper separators
output_path_with_correct_format = 'output/personal_information_with_correct_format.docx'
doc.save(output_path_with_correct_format)





