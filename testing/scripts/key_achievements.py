from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Helper: Apply custom formatting
def apply_custom_formatting(run, formatting):
    if formatting.get("font_style") == "bold":
        run.bold = True
    elif formatting.get("font_style") == "regular":
        run.bold = False
    if "font_size" in formatting:
        run.font.size = Pt(formatting["font_size"])

# Helper: Add line spacing to a paragraph
def apply_line_spacing(paragraph, spacing):
    paragraph.paragraph_format.space_before = Pt(spacing["before"])
    paragraph.paragraph_format.space_after = Pt(spacing["after"])

# Process Key Achievements Section
def process_key_achievements(doc, achievements_data):
    for paragraph in doc.paragraphs:
        if "{{key_achievements}}" in paragraph.text:
            paragraph.text = "Key Achievements"
            paragraph.style = doc.styles["Heading 1"]  # Use a heading style or customize as needed
            apply_line_spacing(paragraph, achievements_data["formatting"]["line_spacing"])

            # Add each achievement
            for entry in achievements_data["entries"]:
                # Add title
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(entry["title"])
                apply_custom_formatting(title_run, achievements_data["formatting"]["title"])
                apply_line_spacing(title_paragraph, achievements_data["formatting"]["line_spacing"])

                # Add description (as an indented bullet)
                description_paragraph = doc.add_paragraph(style="List Bullet")
                description_run = description_paragraph.add_run(entry["description"])
                apply_custom_formatting(description_run, achievements_data["formatting"]["description"])
                apply_line_spacing(description_paragraph, achievements_data["formatting"]["line_spacing"])
            break

# Render Resume Key Achievements Test
def render_resume_key_achievements_test(json_path, template_path, output_path):
    import json
    with open(json_path, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)

    doc = Document(template_path)
    process_key_achievements(doc, data["key_achievements"])
    doc.save(output_path)

# Paths for testing
json_path = "test_data/key_achievements.json"
template_path = "test_templates/key_achievements_template.docx"
output_path = "output/test_key_achievements.docx"

# Run the script
render_resume_key_achievements_test(json_path, template_path, output_path)
