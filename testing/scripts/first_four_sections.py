from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Helper: Apply formatting to a run
def apply_formatting(run, formatting):
    """Applies formatting to a text run."""
    if "font_size" in formatting:
        run.font.size = Pt(formatting["font_size"])
    if "font_style" in formatting:
        if formatting["font_style"] == "bold":
            run.bold = True
        elif formatting["font_style"] == "italic":
            run.italic = True
        elif formatting["font_style"] == "regular":
            run.bold = False
            run.italic = False


# Helper: Strip placeholder and replace in a specific area
def strip_and_replace_placeholder(doc, placeholder, content, formatting=None):
    """Replaces a placeholder with given content in the document."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Remove placeholder and insert new content
            paragraph.text = paragraph.text.replace(placeholder, "")
            if content:
                run = paragraph.add_run(content)
                if formatting:
                    apply_formatting(run, formatting)
            return True  # Exit after replacing the placeholder
    return False  # If placeholder is not found


# Process Personal Information
def process_personal_information(doc, personal_info):
    """Processes and formats personal information section."""
    strip_and_replace_placeholder(doc, "{{name}}", personal_info["name"], personal_info["formatting"]["name"])
    strip_and_replace_placeholder(doc, "{{desired_role}}", personal_info["desired_role"], personal_info["formatting"]["desired_role"])
    contact_info = (
        f"Phone: {personal_info['contact_info']['phone']} | "
        f"Email: {personal_info['contact_info']['email']} | "
        f"LinkedIn: {personal_info['contact_info']['linkedin']} | "
        f"Location: {personal_info['contact_info']['location']}"
    )
    strip_and_replace_placeholder(doc, "{{contact_info}}", contact_info, personal_info["formatting"]["contact_info"])


# Process Summary
def process_summary(doc, summary_data):
    """Processes and formats the summary section."""
    summary_content = f"{summary_data['static_content']} {summary_data['dynamic_content']}"
    strip_and_replace_placeholder(doc, "{{summary}}", summary_content, summary_data["formatting"])


# Process Skills
def process_skills(doc, skills_data):
    """Processes and formats the skills section."""
    skills_content = ""
    for group in skills_data["groups"]:
        skills_content += f"{group['group_name']}: " + ", ".join(group["skills_list"]) + "\n"
    strip_and_replace_placeholder(doc, "{{skills}}", skills_content.strip(), skills_data["formatting"]["skills_list"])


# Process Experience
def process_experience(doc, experience_data):
    """Processes the experience section."""
    # Replace Experience Header
    strip_and_replace_placeholder(doc, "{{experience_header}}", "Experience", {"font_size": 14, "font_style": "bold"})

    # Replace Experience Bullets
    experience_content = ""
    for role in experience_data["roles"]:
        experience_content += f"{role['title']} - {role['company']} ({role['dates']})\n"
        experience_content += f"{role['location']}\n"
        for responsibility in role["responsibilities"]:
            experience_content += f"• {responsibility}\n"
        experience_content += "\n"  # Add spacing between roles

    # Replace the bullets placeholder
    strip_and_replace_placeholder(doc, "{{experience_bullets}}", experience_content.strip(), {"font_size": 11, "font_style": "regular"})


# Render Resume
def render_resume(json_path, template_path, output_path):
    """Generates a tailored resume based on JSON data and a Word template."""
    # Load JSON data
    import json
    with open(json_path, "r", encoding="utf-8") as file:
        data = json.load(file)

    # Load Word template
    doc = Document(template_path)

    # Process each section
    process_personal_information(doc, data["personal_information"])
    process_summary(doc, data["summary"])
    process_skills(doc, data["skills"])
    process_experience(doc, data["experience"])

    # Save the final document
    doc.save(output_path)


# Paths for testing
json_path = "data/tailored_resume.json"  # Update with your JSON file path
template_path = "templates/test_template.docx"  # Update with your template file path
output_path = "output/test_resume_name_role_contact_summary.docx"  # Output file for testing

# Render the resume
render_resume(json_path, template_path, output_path)

