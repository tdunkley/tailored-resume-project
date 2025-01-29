import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Helper: Apply formatting to a paragraph
def apply_formatting(paragraph, formatting):
    if formatting.get("font_size"):
        for run in paragraph.runs:
            run.font.size = Pt(formatting["font_size"])
    if formatting.get("font_style"):
        for run in paragraph.runs:
            if formatting["font_style"] == "bold":
                run.bold = True
            elif formatting["font_style"] == "italic":
                run.italic = True
    if formatting.get("alignment"):
        paragraph.alignment = formatting["alignment"]

# Process education section
def process_education(doc, education_data):
    for paragraph in doc.paragraphs:
        if "{{education}}" in paragraph.text:
            paragraph.text = "Education"
            apply_formatting(paragraph, {"font_size": 14, "font_style": "bold", "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT})
            
            for edu in education_data:
                # Add degree and field
                degree_paragraph = doc.add_paragraph()
                degree_text = f"{edu['degree']}, {edu['field']}"
                degree_run = degree_paragraph.add_run(degree_text)
                degree_run.bold = True
                degree_run.font.size = Pt(12)
                
                # Add institution
                institution_paragraph = doc.add_paragraph()
                institution_paragraph.add_run(edu["institution"]).font.size = Pt(11)
            break

# Render resume with education section
def render_resume_education_test(json_path, template_path, output_path):
    with open(json_path, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)

    doc = Document(template_path)
    process_education(doc, data["education"])
    doc.save(output_path)

# Paths for testing
json_path = "test_data/education.json"  # Ensure this file contains valid education data
template_path = "test_templates/education_template.docx"  # Ensure it contains {{education}}
output_path = "output/test_education.docx"

# Run the script
render_resume_education_test(json_path, template_path, output_path)
