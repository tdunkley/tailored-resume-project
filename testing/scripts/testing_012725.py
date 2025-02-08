from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document Title', 0)

# Add a paragraph
p = doc.add_paragraph('This is a paragraph in the document.')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a run with specific font size
run = p.add_run(' This is a run with a specific font size.')
run.font.size = Pt(12)

# Save the document
doc.save('test.docx')


def example_function():
    """Example function with proper formatting."""
    print("This is an example function.")


class ExampleClass:
    """Example class with proper formatting."""

    def __init__(self):
        """Initialize the example class."""
        self.message = "Hello, World!"

    def display_message(self):
        """Display the message."""
        print(self.message)


def another_function():
    """Another function with proper formatting."""
    print("This is another function.")
