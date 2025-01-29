from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, url, text, color, underline):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(10 * 2))
    rPr.append(sz)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def process_line(paragraph, line):
    if "**" in line:
        bold_parts = line.split("**")
        for i, part in enumerate(bold_parts):
            if i % 2 == 1:  # Odd indices are bold
                run = paragraph.add_run(part.strip())
                run.bold = True
            else:
                paragraph.add_run(part.strip())
    else:
        paragraph.add_run(line)

def render_resume(template_path, data, output_path):
    doc = Document(template_path)
    
    # Set default paragraph spacing
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    processed_sections = set()  # Keep track of processed sections

    for para in doc.paragraphs:
        if "{{header_block}}" in para.text:
            para.text = ""
            
            # Name
            name_run = para.add_run(data["name"].upper())
            name_run.bold = True
            name_run.font.size = Pt(16)
            name_run.font.color.rgb = RGBColor(47, 84, 150)
            para.paragraph_format.space_after = Pt(6)
            para.add_run("\n")
            
            # Role
            role_run = para.add_run(data["desired_role"])
            role_run.font.size = Pt(12)
            role_run.italic = True
            para.paragraph_format.space_after = Pt(6)
            para.add_run("\n")
            
            # Contact info
            contact_parts = data["contact_info"].split("|")
            for i, part in enumerate(contact_parts):
                part = part.strip()
                if "linkedin.com" in part.lower():
                    url = part.split(": ")[1].strip()
                    add_hyperlink(para, url, "LinkedIn Profile", "0000FF", True)
                else:
                    contact_run = para.add_run(part)
                    contact_run.font.size = Pt(10)
                if i < len(contact_parts) - 1:
                    para.add_run(" | ")
            
            para.paragraph_format.space_after = Pt(12)
            para.add_run("\n")
            
        # Process other sections
        for section_key, section_value in data.items():
            if f"{{{{{section_key}}}}}" in para.text and section_key != "header_block":
                if section_key in processed_sections:  # Skip if already processed
                    continue
                    
                para.text = ""
                para.paragraph_format.space_before = Pt(12)
                
                if section_key == 'education':
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_after = Pt(12)
                    
                    for edu_block in section_value:
                        # Process school line
                        process_line(para, edu_block["school"])
                        para.add_run('\n')
                        
                        # Process degree line
                        process_line(para, edu_block["degree"])
                        
                        # Add spacing between blocks (except for last block)
                        if edu_block != section_value[-1]:
                            para.add_run('\n\n')
                            
                    processed_sections.add('education')  # Mark education as processed
                else:
                    # Handle other sections as before
                    for line in section_value.split("\n"):
                        if line.startswith("- "):
                            line = "• " + line[2:]
                            para.paragraph_format.space_after = Pt(3)
                        
                        if "**" in line:
                            bold_parts = line.split("**")
                            for i, part in enumerate(bold_parts):
                                if i % 2 == 1:
                                    run = para.add_run(part.strip())
                                    run.bold = True
                                else:
                                    para.add_run(part.strip())
                        else:
                            para.add_run(line)
                        
                        if line.strip():
                            para.add_run("\n")

                if para.runs and para.runs[-1].text.endswith("\n"):
                    para.runs[-1].text = para.runs[-1].text.rstrip()

    doc.save(output_path)

if __name__ == "__main__":
    template_path = "resume_template.docx"
    output_path = "test_output.docx"
    data = {
        "name": "Troy D. Dunkley",
        "desired_role": "Data Analytics Executive",
        "contact_info": (
            "Phone: +1-770-401-6527 | Email: tdunkley@gmail.com | "
            "LinkedIn: https://www.linkedin.com/in/troy-d-dunkley | Location: Atlanta, GA"
        ),
        "summary": "Visionary Data and Analytics Executive with over 15 years of experience.",
        "skills": (
            "**Technical Skills**: Python, SQL, Power BI, AWS\n"
            "**Leadership Skills**: Team Management, Strategic Planning"
        ),
        "experience": (
            "**Black Nisus, LLC** | **Data Evangelist** | 2020 - Present | Atlanta, GA\n\n"
            "• Developed and implemented advanced BI solutions.\n"
            "• Improved data governance, boosting accuracy by 30%.\n"
            "• Built automated pipelines, reducing manual data processing by 40%.\n\n\n"
        ),
        "education": [
            {
                "school": "**Flatiron School**| Atlanta, GA",
                "degree": "**Certificate of Completion**| Data Science"
            },
            {
                "school": "**Keller Graduate School of Management**| Atlanta, GA",
                "degree": "**Master of Business Administration (MBA)**| Information Systems Management"
            },
            {
                "school": "**Georgia State University**| Atlanta, GA",
                "degree": "**Bachelor of Business Administration (BBA)**| Finance"
            }
        ],
        "key_achievements": (
            "**Spearheaded Digital Transformation Initiative**\n"
            "Orchestrated a company-wide initiative, improving efficiency by 30%.\n\n\n"
            "**Enhanced Data Analytics Capabilities**\n"
            "Built scalable solutions, driving insights for strategic decisions."
        ),
        "projects": (
            "**UA Archive** | 2023 | Remote\n\n"
            "**Goal**: Archive legacy analytics reports for compliance.\n\n"
            "**Responsibilities**: Developed scalable export pipeline using AWS and BigQuery.\n\n"
            "**Results**: Reduced manual effort by 40%.\n\n\n"
        ),
    }

    render_resume(template_path, data, output_path)
