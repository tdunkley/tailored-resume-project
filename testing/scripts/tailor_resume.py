import os
from docx import Document

def tailor_resume(job_description_path, full_cv_path, output_path):
    # Read the job description
    with open(job_description_path, "r") as jd_file:
        job_description = jd_file.read()

    # Open the full CV
    cv_doc = Document(full_cv_path)

    # Analyze and tailor the resume (simplified example)
    tailored_content = []
    for paragraph in cv_doc.paragraphs:
        if any(keyword in paragraph.text.lower() for keyword in job_description.lower().split()):
            tailored_content.append(paragraph.text)

    # Create the tailored resume
    tailored_doc = Document()
    tailored_doc.add_heading("Tailored Resume", level=1)
    for line in tailored_content:
        tailored_doc.add_paragraph(line)

    # Save the tailored resume
    tailored_doc.save(output_path)
    print(f"Tailored resume saved to {output_path}")

if __name__ == "__main__":
    # Define paths
    job_description_path = "job_descriptions/job_description.txt"
    full_cv_path = "resumes/full_cv.docx"
    output_path = "output/tailored_resume.docx"

    # Ensure directories exist
    if not os.path.exists("output"):
        os.makedirs("output")

    # Tailor the resume
    tailor_resume(job_description_path, full_cv_path, output_path)
