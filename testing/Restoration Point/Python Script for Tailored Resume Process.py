import json
from docx import Document


def load_json(file_path):
    with open(file_path, 'r') as file:
        return json.load(file)


def tailor_resume(job_description, master_resume, full_cv, output_path):
    # Load data
    job_description_data = load_json(job_description)
    master_resume_data = load_json(master_resume)
    full_cv_data = load_json(full_cv)

    # Open the Word document for editing
    doc = Document()

    # Personal Information
    personal_info = master_resume_data['personal_information']
    doc.add_heading(personal_info['name'], level=1)
    doc.add_paragraph(
        f"{personal_info['contact_info']['phone']} | {personal_info['contact_info']['email']} | "
        f"{personal_info['contact_info']['location']} | {personal_info['contact_info']['linkedin']}"
    )

    # Summary
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(master_resume_data['summary']['content'])

    # Skills
    doc.add_heading('Skills', level=2)
    skills_section = master_resume_data['skills']
    for skill_category in ['static', 'dynamic']:
        if skill_category in skills_section:
            doc.add_paragraph(f"{skill_category.capitalize()} Skills:")
            for skill in skills_section[skill_category]:
                doc.add_paragraph(f"• {skill}", style='List Bullet')

    # Experience
    doc.add_heading('Experience', level=2)
    for role in master_resume_data['experience']['roles']:
        role_section = doc.add_paragraph()
        role_section.add_run(f"{role['title']} - {role['company']}").bold = True
        role_section.add_run(f"\n{role['location']} | {role['dates']}")
        for responsibility in role['responsibilities']:
            doc.add_paragraph(f"• {responsibility}", style='List Bullet')

    # Education
    doc.add_heading('Education', level=2)
    for entry in master_resume_data['education']['entries']:
        edu_section = doc.add_paragraph()
        edu_section.add_run(f"{entry['degree']} - {entry['school']}").bold = True
        edu_section.add_run(f"\n{entry['location']} | {entry['focus']}")

    # Key Achievements
    doc.add_heading('Key Achievements', level=2)
    for achievement in master_resume_data['key_achievements']['entries']:
        achievement_section = doc.add_paragraph()
        achievement_section.add_run(achievement['title']).bold = True
        achievement_section.add_run(f"\n{achievement['description']}")

    # Projects
    doc.add_heading('Projects', level=2)
    for project in master_resume_data['projects']['entries']:
        project_section = doc.add_paragraph()
        project_section.add_run(project['title']).bold = True
        project_section.add_run(f"\nGoal: {project['goal']}")
        project_section.add_run(f"\nResponsibilities: {project['responsibilities']}")
        project_section.add_run(f"\nResults: {project['results']}")

    # Save the tailored resume
    doc.save(output_path)
    print(f"Tailored resume saved to {output_path}")


if __name__ == "__main__":
    # Paths for input and output files
    job_description_file = "job_description.json"
    master_resume_file = "master_resume.json"
    full_cv_file = "full_cv.json"
    tailored_resume_file = "tailored_resume.docx"

    # Tailor the resume
    tailor_resume(job_description_file, master_resume_file, full_cv_file, tailored_resume_file)
