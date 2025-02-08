from docx import Document
import os
from tracker import log_event
from config_core import OUTPUT_DIR, RESUME_FILE, load_json

def apply_formatting(paragraph, text, style):
    """Apply formatting to a paragraph using a built-in style."""
    run = paragraph.add_run(text)
    paragraph.style = style

def generate_full_resume(resume_data, config_data, output_file):
    """Generate a full resume document."""
    # Create a new document from scratch
    document = Document()
    for section_name, section_content in resume_data.get("sections", {}).items():
        # Add section header using built-in style
        header = document.add_paragraph()
        apply_formatting(header, section_name.upper(), 'Heading1')
        
        for key, value in section_content.items():
            # Add sub-header using built-in style
            sub_header = document.add_paragraph()
            apply_formatting(sub_header, key, 'Heading2')
            
            # Add content using normal style
            paragraph = document.add_paragraph()
            apply_formatting(paragraph, str(value), 'Normal')
    
    document.save(output_file)
    log_event("FULL_RESUME", "SUCCESS", f"Resume generated: {output_file}")

if __name__ == "__main__":
    resume_data = load_json(RESUME_FILE)
    config_data = load_json("config.json")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, "full_resume.docx")
    generate_full_resume(resume_data, config_data, output_path)
    print(f"Full resume generated at {output_path}")
