import os
import json
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Setup logging
logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")

def apply_dynamic_formatting(run, formatting):
    """Apply dynamic formatting to a run based on the configuration."""
    run.font.name = formatting.get("font", "Arial")
    run.font.size = Pt(formatting.get("size", 12))
    run.font.bold = formatting.get("bold", False)
    run.font.color.rgb = RGBColor.from_string(formatting.get("color", "000080").replace("#", ""))
    return run

def add_personal_information(document, personal_info, formatting):
    """Add the personal information section to the document without a section header."""
    logging.info("Adding personal information")
    logging.info(f"Personal information data: {personal_info}")

    # Add name and desired role
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(personal_info["name"])
    apply_dynamic_formatting(name_run, formatting.get("name", {}))
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    role_paragraph = document.add_paragraph()
    role_run = role_paragraph.add_run(personal_info["desired_role"])
    apply_dynamic_formatting(role_run, formatting.get("desired_role", {}))
    role_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add contact info in sentence format
    contact_info = personal_info["contact_info"]
    contact_info_text = f"{contact_info['phone']} | {contact_info['email']} | {contact_info['location']} | {contact_info['linkedin']}"
    contact_paragraph = document.add_paragraph()
    contact_run = contact_paragraph.add_run(contact_info_text)
    apply_dynamic_formatting(contact_run, formatting.get("contact_info", {}))
    contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def main():
    """Main function to test the personal information section formatting."""
    logging.info("Starting the document generation process")
    document = Document()

    # Load configuration
    config_path = os.path.join(os.path.dirname(__file__), "config.json")
    with open(config_path, 'r') as file:
        config = json.load(file)

    # Sample personal information data
    personal_info = {
        "name": "Troy D. Dunkley",
        "desired_role": "Senior Data Architect",
        "contact_info": {
            "phone": "+1-770-401-6527",
            "email": "tdunkley@gmail.com",
            "location": "Atlanta, GA",
            "linkedin": "https://www.linkedin.com/in/troy-d-dunkley"
        }
    }

    # Add personal information to the document
    add_personal_information(document, personal_info, config["formatting"])

    # Save the document
    output_dir = os.path.join(os.path.dirname(__file__), "output")
    if not os.path.exists(output_dir):
        logging.info(f"Creating output directory: {output_dir}")
        os.makedirs(output_dir)
    output_path = os.path.join(output_dir, "test_formatting.docx")
    document.save(output_path)
    logging.info(f"Test document generated and saved to {output_path}")

if __name__ == "__main__":
    main()
