import os
import json
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def apply_dynamic_formatting(run, formatting):
    """Apply dynamic formatting to a run based on the configuration."""
    logging.debug(f"Applying formatting: {formatting}")
    run.font.name = formatting.get("font", "Arial")
    run.font.size = Pt(formatting.get("size", 12))  # Default size is 12 if not specified
    run.font.bold = formatting.get("bold", False)
    if "color" in formatting:
        try:
            hex_color = formatting.get("color", "000000").replace("#", "")
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
            logging.debug(f"Color set to RGB({r}, {g}, {b})")
        except Exception as e:
            logging.error(f"Error setting color: {e}")

    return run

def add_hyperlink(paragraph, url, text, formatting):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = paragraph.add_run(text)
    apply_dynamic_formatting(new_run, formatting)
    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

def add_personal_information(document, personal_info, formatting):
    """Add the personal information section to the document."""
    logging.info("Adding personal information")
    
    # Add Name as a standalone paragraph
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(personal_info["name"])
    
    # Apply formatting for "name" from config.json
    apply_dynamic_formatting(name_run, formatting.get("name", {}))  # Use "name" formatting
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Desired Role
    role_paragraph = document.add_paragraph()
    role_run = role_paragraph.add_run(personal_info["desired_role"])
    apply_dynamic_formatting(role_run, formatting.get("desired_role", {}))
    role_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Contact Info
    contact_info = personal_info["contact_info"]
    contact_info_text = f"{contact_info['phone']} | {contact_info['email']} | {contact_info['location']} | "
    
    contact_paragraph = document.add_paragraph()
    
    contact_run = contact_paragraph.add_run(contact_info_text)
    
     # Apply Contact Info Formatting
    apply_dynamic_formatting(contact_run, formatting.get("contact_info", {}))

     # Add LinkedIn Hyperlink
    add_hyperlink(contact_paragraph, contact_info['linkedin'], contact_info['linkedin'], formatting.get("contact_info", {}))

def main():
    """Main function to test the personal information section formatting."""
    document = Document()

    # Load configuration
    config_path = os.path.join(os.path.dirname(__file__), "config.json")
    with open(config_path, 'r') as file:
        config = json.load(file)

    # Sample personal information data
    personal_info = {
        "name": "Troy D. Dunkley",
        "desired_role": "Senior Data Architect",
        "contact_info": {
            "phone": "+1-770-401-6527",
            "email": "tdunkley@gmail.com",
            "location": "Atlanta, GA",
            "linkedin": "https://www.linkedin.com/in/troy-d-dunkley"
        }
    }

    # Add personal information to the document
    add_personal_information(document, personal_info, config["formatting"])

    # Ensure output directory exists
    output_dir = os.path.join(os.path.dirname(__file__), "output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the document to the output directory
    output_path = os.path.join(output_dir, "test_personal_information_v2.docx")
    try:
        document.save(output_path)
        logging.info(f"Test document generated and saved to {output_path}")
        print(f"Test document generated and saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save the document: {e}")
        print(f"Failed to save the document: {e}")

if __name__ == "__main__":
    main()
